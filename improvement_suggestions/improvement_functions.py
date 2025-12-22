from match_evaluation.agent_state import AgentState
from extracting_data.description_schemas import CVDescription, JobDescription
from improvement_suggestions.improvement_state import CVRewriteState
from improvement_suggestions.improvement_prompts import (
    MAIN_UPDATE_CV_PROMPT,
    CV_FEEDBACK_UPDATE_PROMPT,
)
from improvement_suggestions.improvement_output_schema import UpdatedCvResult
from improvement_suggestions.improvement_consts import (
    SAVING_FONT,
    UPDATED_CV_NAME,
    AMOUNT_FEEDBACK_ROUNDS,
)
import os
from docx import Document
from docx.shared import Pt
import re
from pathlib import Path
from typing import Any


def render_experience(experiences):
    """
    Renders experience entries into a flat, readable text block
    suitable for LLM prompt injection.
    """

    blocks = []

    for exp in experiences:
        header = f"{exp.title} | {exp.company} | {exp.start_date} - {exp.end_date}"
        blocks.append(header)

        if getattr(exp, "responsibilities", None):
            for r in exp.responsibilities:
                blocks.append(f"- {r}")

        if getattr(exp, "quantifiable_achievements", None):
            for q in exp.quantifiable_achievements:
                blocks.append(f"- {q}")

        blocks.append("")

    return "\n".join(blocks).strip()


def render_education(education_list):
    lines = []
    for e in education_list:
        parts = []
        if e.certification:
            parts.append(e.certification)
        if e.field:
            parts.append(f"in {e.field}")
        if e.institution:
            parts.append(e.institution)
        if e.graduation_year:
            parts.append(f"({e.graduation_year})")

        if parts:
            lines.append(" ".join(parts))

    return "\n".join(lines)


def add_runs_with_formatting(paragraph, text):
    """
    Handles **bold** and *italic* inline formatting.
    """
    tokens = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)

    for token in tokens:
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def prompt_user_to_cv_rewrite(
    state: AgentState,
):
    continue_decision = input(
        "Do you want to update your CV to match profile better? Answer yes or no\n"
    )
    answered_understood = False
    while answered_understood == False:
        if any([continue_decision.lower() in answer for answer in ["y", "yes"]]):
            answered_understood = True
            return "rewrite"
        elif any([continue_decision.lower() in answer for answer in ["n", "no"]]):
            answered_understood = True
            return "finish"
        else:
            continue_decision = input(
                'Please answer "yes" or "no", to answer whether you want to update your CV to match profile better?\n'
            )


def prompt_user_satisfaction(state: CVRewriteState):
    print("\nThis is the new version of your CV:")
    print(state.updated_cv_text)
    user_satisfaction = input(
        "Would you like to change anything in the created CV? Answer yes or no\n"
    )
    answered_understood = False
    while answered_understood == False:
        if any([user_satisfaction.lower() in answer for answer in ["y", "yes"]]):
            answered_understood = True
            return "change needed"
        elif any([user_satisfaction.lower() in answer for answer in ["n", "no"]]):
            answered_understood = True
            return "no change"
        else:
            user_satisfaction = input(
                'Please answer "yes" or "no", to answer whether you want to continue to updating your CV to match profile better?\n'
            )


def receive_user_feedback(state: CVRewriteState, llm):
    input = input("'nPlease share your thoughts on what needs to be changed\n")
    return {"user_feedback": input}


def _state_get(state, key, default=None):
    if isinstance(state, dict):
        return state.get(key, default)
    return getattr(state, key, default)


def _coerce_to_model(value, model_cls):
    """
    Ensure we pass actual model instances (or dicts convertible) to CVRewriteState.
    Handles BaseModel-like objects, dicts, or already-correct instances.
    """
    if value is None:
        return None
    if isinstance(value, model_cls):
        return value
    if hasattr(value, "model_dump"):
        try:
            value = value.model_dump()
        except Exception:
            pass
    if isinstance(value, dict):
        try:
            return model_cls.model_validate(value)
        except Exception:
            return value
    return value


def _invoke_updated_cv(prompt: str, llm: Any) -> dict:
    """
    Try structured output; if the model returns an unexpected shape, fall back to raw content.
    """
    try:
        result = llm.with_structured_output(UpdatedCvResult).invoke(prompt)
        if hasattr(result, "updated_cv_text"):
            return {"updated_cv_text": result.updated_cv_text}
    except Exception:
        pass

    raw = llm.invoke(prompt)
    updated_cv_text = (
        getattr(raw, "content", None) or getattr(raw, "text", None) or str(raw)
    )
    return {"updated_cv_text": updated_cv_text}


def create_rewrite_state(state: AgentState, llm) -> CVRewriteState:
    """
    Extract only the essential information needed for CV rewriting.
    This reduces token usage and focuses the LLM on relevant data.
    """
    path_to_cv = _state_get(state, "path_to_cv")
    original_folder = (
        "/".join(path_to_cv.split("/")[:-1])
        if path_to_cv
        else str((Path.cwd() / "outputs").resolve())
    )
    output_dir = original_folder
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    original_cv = _coerce_to_model(_state_get(state, "cv"), CVDescription)
    target_job = _coerce_to_model(_state_get(state, "job"), JobDescription)
    return CVRewriteState(
        # Core documents
        original_cv_folder_path=original_folder,
        output_dir=output_dir,
        original_cv=original_cv,
        target_job=target_job,
        original_cv_text=_state_get(state, "cv_description_text"),
        # Skills optimization
        matched_skills=_state_get(state, "skills_match").matched_items,
        partial_skill_matches=_state_get(state, "skills_match").partial_matches,
        matched_keywords=_state_get(state, "keyword_match").matched_keywords,
        missing_keywords=_state_get(state, "keyword_match").missing_keywords,
        # Requirements
        must_haves_satisfied=_state_get(
            state, "requirements_coverage"
        ).must_have_satisfied,
        must_haves_missing=_state_get(state, "requirements_coverage").must_have_missing,
        nice_to_haves_satisfied=_state_get(
            state, "requirements_coverage"
        ).nice_to_have_satisfied,
        # Experience context
        recent_relevant_experience=_state_get(
            state, "recency_relevance"
        ).recent_relevant_experience,
        matched_domains=_state_get(state, "domain_match").matched_items,
        transferable_domains=_state_get(state, "domain_match").transferable_experience,
        # Seniority
        candidate_level=_state_get(state, "seniority_match").candidate_level,
        required_level=_state_get(state, "seniority_match").required_level,
        title_alignment=_state_get(state, "seniority_match").title_alignment,
        # Strategic priorities
        top_strengths=_state_get(state, "final_scoring").strengths[:5],
        key_weaknesses=_state_get(state, "final_scoring").weaknesses[:3],
        red_flags=_state_get(state, "all_red_flags"),
        keyword_frequency_targets=_state_get(state, "keyword_match").keyword_frequency,
        focus_areas=_state_get(state, "final_scoring").focus_areas or [],
    )


def rewrite_cv_initial(state: CVRewriteState, llm):
    updating_cv_prompt = MAIN_UPDATE_CV_PROMPT.format(
        full_name=state.original_cv.full_name,
        current_title=state.original_cv.current_title,
        original_cv_text=state.original_cv_text,
        total_years_experience=state.original_cv.total_years_experience,
        cv_domains=", ".join(state.original_cv.domains),
        technical_skills="\n".join(
            f"- {s.name}" for s in state.original_cv.technical_skills
        ),
        soft_skills="\n".join(f"- {s.name}" for s in state.original_cv.soft_skills),
        experience_history=render_experience(state.original_cv.experience),
        projects="\n".join(p.name for p in state.original_cv.projects),
        education=render_education(state.original_cv.education),
        certifications="\n".join(state.original_cv.certifications),
        languages="\n".join(state.original_cv.spoken_languages),
        job_title=state.target_job.job_title,
        company=state.target_job.company,
        required_years_experience=state.target_job.required_years_experience,
        required_seniority=state.target_job.required_seniority,
        required_domains=", ".join(state.target_job.required_domains),
        job_responsibilities="\n".join(state.target_job.responsibilities),
        job_required_skills="\n".join(
            [skill.name for skill in state.target_job.required_technical_skills]
        ),
        job_nice_to_have_skills="\n".join(
            [skill.name for skill in state.target_job.nice_to_have_skills]
        ),
        job_critical_keywords="\n".join(state.target_job.critical_keywords),
        job_role_summary=state.target_job.role_summary,
        matched_skills="\n".join(state.matched_skills),
        partial_skill_matches="\n".join(state.partial_skill_matches),
        matched_keywords="\n".join(state.matched_keywords),
        missing_keywords="\n".join(state.missing_keywords),
        keyword_frequency_targets=str(state.keyword_frequency_targets),
        must_haves_satisfied="\n".join(state.must_haves_satisfied),
        must_haves_missing="\n".join(state.must_haves_missing),
        nice_to_haves_satisfied="\n".join(state.nice_to_haves_satisfied),
        recent_relevant_experience="\n".join(state.recent_relevant_experience),
        matched_domains="\n".join(state.matched_domains),
        transferable_domains="\n".join(state.transferable_domains),
        candidate_level=state.candidate_level,
        required_level=state.required_level,
        title_alignment=state.title_alignment,
        top_strengths="\n".join(state.top_strengths),
        key_weaknesses="\n".join(state.key_weaknesses),
        red_flags="\n".join(state.red_flags),
        focus_areas="\n".join(state.focus_areas),
    )
    return _invoke_updated_cv(updating_cv_prompt, llm)


def rewrite_cv_with_feedback(state: CVRewriteState, llm):
    updating_cv_with_feedback_prompt = CV_FEEDBACK_UPDATE_PROMPT.format(
        user_feedback=state.user_feedback,
        updated_cv_text=state.updated_cv_text,
        full_name=state.original_cv.full_name,
        total_years_experience=state.original_cv.total_years_experience,
        original_experience_excerpt="\n".join(
            f"- {exp.title} at {exp.company}: {', '.join(exp.responsibilities[:2])}"
            for exp in state.original_cv.experience
        ),
        original_skills=", ".join(s.name for s in state.original_cv.technical_skills),
        original_projects="\n".join(
            f"- {p.name}: {p.project_description}" for p in state.original_cv.projects
        ),
        job_title=state.target_job.job_title,
        company=state.target_job.company,
        matched_keywords="\n".join(state.matched_keywords[:10]),
        must_haves_satisfied="\n".join(state.must_haves_satisfied[:5]),
    )
    result = _invoke_updated_cv(updating_cv_with_feedback_prompt, llm)
    feedback_round = state.feedback_round + 1
    result["feedback_round"] = feedback_round
    return result


def markdown_to_docx(state: CVRewriteState, llm) -> None:
    doc = Document()
    markdown_text = state.updated_cv_text
    normal_style = doc.styles["Normal"]
    normal_style.font.name = SAVING_FONT
    normal_style.font.size = Pt(11)

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        # ---- Headings ----
        if line.startswith("### "):
            p = doc.add_heading(level=3)
            run = p.add_run(line[4:])
            run.font.name = SAVING_FONT
            run.font.size = Pt(12)
            run.bold = True

        elif line.startswith("## "):
            p = doc.add_heading(level=2)
            run = p.add_run(line[3:])
            run.font.name = SAVING_FONT
            run.font.size = Pt(13)
            run.bold = True

        elif line.startswith("# "):
            p = doc.add_heading(level=1)
            run = p.add_run(line[2:])
            run.font.name = SAVING_FONT
            run.font.size = Pt(14)
            run.bold = True

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_formatting(p, line[2:])

        elif re.match(r"\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            content = re.sub(r"^\d+\.\s+", "", line)
            add_runs_with_formatting(p, content)

        elif line == "":
            continue

        # ---- Normal paragraph ----
        else:
            p = doc.add_paragraph()
            add_runs_with_formatting(p, line)

    output_dir = state.output_dir or state.original_cv_folder_path
    Path(output_dir).mkdir(parents=True, exist_ok=True)
    save_path = os.path.join(output_dir, UPDATED_CV_NAME)
    doc.save(save_path)
    return {"docx_path": save_path}
